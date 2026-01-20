from docx import Document
from io import BytesIO
from datetime import datetime

def plan_to_docx(title: str, goal: str, constraints: dict, plan_text: str):
    doc = Document()

    # Title
    doc.add_heading(title, level=1)

    # Metadata
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Goal:\n{goal}")

    doc.add_heading("Constraints", level=2)
    for k, v in constraints.items():
        doc.add_paragraph(f"- {k.replace('_', ' ').title()}: {v}")

    doc.add_heading("Plan", level=2)
    for line in plan_text.split("\n"):
        doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
